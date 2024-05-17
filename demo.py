#!/usr/bin/env python
import sys
import json
import numpy as np
from docx import Document
import matplotlib.pyplot as plt

path = sys.argv[1]


def read_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()


# Function to calculate the sum of field_a.value and field_b.value


def calculate_sum(input_json):
    field_a_value = input_json["field_a"]["value"]
    field_b_value = input_json["field_b"]["value"]
    sum_value = field_a_value + field_b_value
    return sum_value


def generate_line_chart_data(num_points, lang):
    num_points = int(num_points)
    if num_points < 0:
        num_points = 1

    x_values = np.linspace(0, 10, num_points)
    y_values1 = np.sin(x_values)
    y_values2 = np.cos(x_values)

    if lang == 'lv':
        label = 'Līnijas grafiks no HPC'
    else:
        label = 'Line chart from HPC'

    line_chart_data = {
        'type': 'line',
        'data': {
            'labels': list(x_values),
            'datasets': [
                {
                    'label': 'Sin',
                    'data': list(y_values1),
                },
                {
                    'label': 'Cos',
                    'data': list(y_values2),
                }
            ]
        },
        'options': {
            "plugins": {
                "legend": {
                    "position": 'top',
                },
                "title": {
                    "display": "true",
                    "text": label
                }
            },
            "scales": {
                "x": {
                    "type": "linear",
                    "position": "bottom"
                },
                "y": {
                    "type": "linear"
                }
            }
        }
    }
    return line_chart_data


def generate_test_msword(input_json, value, lang):
    text = read_file(path + input_json["field_user_text"]["filename"])

    a = input_json["field_a"]["value"]
    b = input_json["field_b"]["value"]

    doc = Document()
    # Add a paragraph with the sample text

    if lang == "lv":
        doc.add_paragraph("Mēs veicāc summas aprēķiņu: " + str(a) + "+" + str(b) + "=" + str(value))
        doc.add_paragraph("Jūs iesniedzāt šādu teksta dokumentu: ")
        doc.add_paragraph(text)
        doc.save(path + "/output/files/my_word_lv.docx")
        return "/output/files/my_word_lv.docx"
    else:
        doc.add_paragraph("We calculated a sum: " + str(a) + "+" + str(b) + "=" + str(value))
        doc.add_paragraph("You have provided the text document below: ")
        doc.add_paragraph(text)
        doc.save(path + "/output/files/my_word_en.docx")
        return "/output/files/my_word_en.docx"


def generate_test_dt(input_json, lang):
    if lang == 'lv':
        oper_label = "Operācija"
        res_label = 'Rezultāts'
        caption = "Aprēķini"
    else:
        oper_label = "Operation"
        res_label = "Result"
        caption = "Calculations"

    dt = {
        'caption': caption,
        'columns': [
            {'title': 'A'},
            {'title': oper_label},
            {'title': 'B'},
            {'title': res_label}
        ],
        'data': [
            [input_json["field_a"]["value"], '+', input_json["field_b"]["value"],
             input_json["field_a"]["value"] + input_json["field_bs"]["value"]],
            [input_json["field_a"]["value"], '-', input_json["field_b"]["value"],
             input_json["field_a"]["value"] - input_json["field_bs"]["value"]],
            [input_json["field_a"]["value"], '*', input_json["field_b"]["value"],
             input_json["field_a"]["value"] * input_json["field_bs"]["value"]],
            [input_json["field_a"]["value"], '/', input_json["field_b"]["value"],
             round(input_json["field_a"]["value"] / input_json["field_bs"]["value"] * 100) / 100],
        ]
    }
    return dt


def generate_test_chart(input_json, sum, lang):
    a = input_json["field_a"]["value"]
    b = input_json["field_b"]["value"]

    # Data for the bars
    categories = ['a', 'b', 'y=a+b']
    values = [a, b, sum]

    # Create a bar chart
    plt.bar(categories, values)

    if lang == 'lv':
        # Add title and labels
        plt.title('Joslu diagrammas piemērs')
        plt.xlabel('Mainīgie')
        plt.ylabel('Vērtības')
    else:
        # Add title and labels
        lang = 'en'
        plt.title('A bar chart example')
        plt.xlabel('Variables')
        plt.ylabel('Values')

    # Save the chart as a PNG file
    plt.savefig(path + "/output/files/" + lang + "_chart.png")
    return "/output/files/" + lang + "_chart.png"


# Read input JSON file
with open(path + "/input/input.json", "r") as f:
    input_data = json.load(f)

# Calculate sum
sum_value = calculate_sum(input_data)

# Create output JSON object
output_data = {
    "values": {
        "sum": sum_value
    },
    "datatables": {
        "tableLv": generate_test_dt(input_data, 'lv'),
        "tableEn": generate_test_dt(input_data, 'en')
    },
    "chartjs": {
        "myLine_lv": generate_line_chart_data(sum_value, 'lv'),
        "myLine_en": generate_line_chart_data(sum_value, 'en')
    },
    "files": {
        "MyWord_lv": generate_test_msword(input_data, sum_value, 'lv'),
        "MyWord_en": generate_test_msword(input_data, sum_value, 'en'),
        "MyChart_lv": generate_test_chart(input_data, sum_value, 'lv'),
        "MyChart_en": generate_test_chart(input_data, sum_value, 'en')
    }
}

# Write output JSON to STDOUT
print(json.dumps(output_data))
